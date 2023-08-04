from docx import Document
import sys

def load_document(file_path):
    return Document(file_path)

"""
Select column in the docx 
Replace with the column_index of the column you want (0-based index)
    for cell_text in selected_column:
        print("cell:")
        print(cell_text)
"""
def select_column(doc, column_index):
    selected_column = []

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) > column_index:
                cell = row.cells[column_index]
                selected_column.append(cell.text)

    return selected_column


def return_paragraphs(doc):
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():  # Ignore empty paragraphs
            paragraphs.append(para.text)
    return paragraphs


def skip_empty_paragraph(para):
    return para.text.strip()


def process_paragraph(doc):
    for para in doc.paragraphs:
        if skip_empty_paragraph(para):  # Ignore empty paragraphs
            para.text = f"{para.text} \n{text_to_add}"
    #         paragraphs.append(para.text)
    # return paragraphs

if __name__ == "__main__":
    sys.stdout.reconfigure(encoding='utf-8')
    doc = load_document('/documents/duizhao_xml.docx')

    text_to_add = "TEXT TO ADD"

    process_paragraph(doc)

    # for idx, para in enumerate(paragraphs, 1):
    #     print(f"Paragraph {idx}: {para}")

    # Save the modified content back to the DOCX file
    doc.save('/Users/bilibala/Study/SLA-TRANS/documents/duizhao_xml_modified.docx')


