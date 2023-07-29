from docx import Document
from docx.oxml import parse_xml
from tools.utils import find_key_containing_val
from tools.colors import print_colored_block
import sys


def load_document(file_path):
    return Document(file_path)


"""
    Select column in the docx 
    Replace with the column_index of the column you want (0-based index)
        for cell_text in selected_column:
            print("cell:")
            print(cell_text)
    Get all the cells list in the selected column
    Return a list of dictionary {'object':cell, 'index': (table_index, row_index, column_index)}
"""
def get_cell_dict_in_selected_column(doc, col_index=0) -> dict:
    cell_dict_in_selected_column = []
    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            if len(row.cells) > col_index:
                cell = row.cells[col_index]
                cell_dict_in_selected_column.append({'object':cell, 'index': (table_index, row_index, col_index)})
    return cell_dict_in_selected_column


def find_cell_by_index(doc, index:tuple):
    table_index, row_index, col_index = index
    try:
        table = doc.tables[table_index]
        cell = table.cell(row_index, col_index)
        return {'object':cell, 'index': (table_index, row_index, col_index)}
    except IndexError:
        return None


def copy_font(source_font, target_font):
    target_font.name = source_font.name
    target_font.size = source_font.size
    target_font.bold = source_font.bold
    target_font.italic = source_font.italic
    target_font.underline = source_font.underline
    target_font.color.rgb = source_font.color.rgb


def copy_paragraph_format(source_format, target_format):
    target_format.left_indent = source_format.left_indent
    target_format.right_indent = source_format.right_indent
    target_format.space_before = source_format.space_before
    target_format.space_after = source_format.space_after
    target_format.line_spacing = source_format.line_spacing
    target_format.alignment = source_format.alignment


# Function to copy cell content and format from source_cell to target_cell
def copy_cell_format(source_cell, target_cell, new_content):
    target_cell.paragraphs.clear()

    for source_paragraph in source_cell.paragraphs:
        target_paragraph = target_cell.add_paragraph()
        # target_paragraph.text = source_paragraph.text

        # Copy the paragraph format
        copy_paragraph_format(source_paragraph.paragraph_format, target_paragraph.paragraph_format)

        # Copy the runs and their formatting
        for source_run in source_paragraph.runs:
            target_run = target_paragraph.add_run(source_run.text)
            copy_font(source_run.font, target_run.font)


def edit_cell_content(doc, index:tuple, new_content:str):
    table_index, row_index, col_index = index
    try:
        table = doc.tables[table_index]
        cell = table.cell(row_index, col_index)
        new_cell = table.cell(row_index, col_index)
        # copy_cell_format(cell, new_cell)
        new_cell.text = new_content
        return True
    except IndexError:
        return False


def delete_cell_content(doc, index:tuple):
    table_index, row_index, col_index = index
    try:
        table = doc.tables[table_index]
        cell = table.cell(row_index, col_index)
        cell.text = ""
        return True
    except IndexError:
        return False


"""
    Get the background HEX shading color from cell
    Returns HEX string value
"""
def get_cell_shading_color(cell):
    shading = cell._element.xpath(".//w:shd")
    if shading and shading[0].attrib:
        color_dict = shading[0].attrib
        key_list = list(color_dict.keys())
        find_key_containing_value = find_key_containing_val(key_list, 'fill')
        color_str = color_dict[find_key_containing_value]
        return color_str
    # white color, no sharding, return "FFFFFF"
    return "FFFFFF"


"""
    Check the unique colors appeared in docs table and Print each colors out
    Returns a set of unique colors 
"""
def check_table_colors(doc) -> set:
    table_colors = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                color = get_cell_shading_color(cell)
                if color:
                    table_colors.add(color)
    for color in table_colors:
        print_colored_block(color)
    return table_colors


def get_cell_text(cell):
    return cell.text


def print_cell_color(cell):
    color = get_cell_shading_color(cell)
    if color:
        print_colored_block(color)
    # else:
    #     print_colored_block("FFFFFF")


def if_color_match_the_cell(cell, color_to_be_checked):
    cell_color = get_cell_shading_color(cell)
    return True if cell_color == color_to_be_checked else False


"""
    Iterate through each cells list
    Print the texts in each cell, and print the color in each cell
"""
def column_processor(doc, column_index_to_select=0):
    cells_in_selected_column = get_cell_dict_in_selected_column(doc, column_index_to_select)
    for cell_dict in cells_in_selected_column:
        print("-" * 50)
        print(cell_dict['object'].text)
        print_cell_color(cell_dict['object'])


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding='utf-8')
    doc_path = '/Users/bilibala/Study/SLA-TRANS/documents/test1.docx'
    doc = load_document(doc_path)

    print("\nINFO: column_processor(doc, 3)")
    column_processor(doc, 3)

    print("\nINFO: check_table_colors(doc)")
    table_colors = check_table_colors(doc)

    print("\nINFO: if_color_match_the_cell")
    for cell_dict in get_cell_dict_in_selected_column(doc, 3):
        if if_color_match_the_cell(cell_dict['object'], "808080"):
            print(get_cell_text(cell_dict['object']))
            print(cell_dict['index'])

    print("\nINFO: find_cell_by_index")
    val_dict = find_cell_by_index(doc, (0, 6, 3))
    if val_dict:
        print(get_cell_text(val_dict['object']))

    print("\nINFO: edit_cell_content")

    content = "test abc content\n"
    if edit_cell_content(doc, (0, 1, 0), content):
        print("Cell content modified successfully.")
    else:
        print("Failed to modify cell content.")
    doc.save(doc_path)

